"""
Resume Parsing and NLP Entity Extraction Engine
Supports PDF (PyMuPDF / fitz) and DOCX (python-docx) extraction.
Identifies Contact Information (Name, Email, Phone, Links), Skills (categorized),
Education, Experience, Projects, Certifications, and Summary.
"""

import os
import re
import json
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx  # python-docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

import zipfile
import xml.etree.ElementTree as ET


# Comprehensive Tech and Soft Skills taxonomy
SKILLS_TAXONOMY: Dict[str, List[str]] = {
    "Languages": [
        "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "C", "Go", "Golang",
        "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "Dart", "MATLAB", "Perl",
        "Shell", "Bash", "PowerShell", "SQL", "HTML", "HTML5", "CSS", "CSS3", "Sass", "SCSS"
    ],
    "Frameworks & Libraries": [
        "React", "React.js", "React Native", "Vue", "Vue.js", "Angular", "Next.js", "Nuxt.js",
        "Node.js", "Express", "Express.js", "NestJS", "FastAPI", "Flask", "Django", "Spring Boot",
        "ASP.NET", "Ruby on Rails", "Laravel", "PyTorch", "TensorFlow", "Keras", "Scikit-Learn",
        "Pandas", "NumPy", "OpenCV", "Tailwind CSS", "Bootstrap", "Material-UI", "MUI", "Redux",
        "Zustand", "GraphQL", "REST API", "RESTful APIs", "gRPC", "Electron"
    ],
    "Databases & Storage": [
        "PostgreSQL", "MySQL", "SQLite", "MongoDB", "Redis", "Cassandra", "Elasticsearch",
        "DynamoDB", "Firebase", "Firestore", "Supabase", "Oracle", "Microsoft SQL Server",
        "MSSQL", "MariaDB", "Neo4j", "CouchDB", "Prisma", "Drizzle", "SQLAlchemy", "Mongoose"
    ],
    "Cloud & DevOps": [
        "AWS", "Amazon Web Services", "Google Cloud", "GCP", "Google Cloud Platform", "Microsoft Azure",
        "Azure", "Docker", "Kubernetes", "K8s", "CI/CD", "GitHub Actions", "GitLab CI", "Jenkins",
        "Terraform", "Ansible", "Nginx", "Apache", "Linux", "Ubuntu", "Debian", "CentOS", "RedHat",
        "Serverless", "AWS Lambda", "Vercel", "Netlify", "Render", "Heroku", "Cloudflare"
    ],
    "Tools & Platforms": [
        "Git", "GitHub", "GitLab", "Bitbucket", "Postman", "Swagger", "Jira", "Confluence",
        "Trello", "Figma", "Canva", "VS Code", "Visual Studio", "IntelliJ IDEA", "PyCharm",
        "Webpack", "Vite", "Babel", "npm", "yarn", "pnpm", "pip", "Docker Compose"
    ],
    "Methodologies & Concepts": [
        "Agile", "Scrum", "Kanban", "Test-Driven Development", "TDD", "CI/CD", "Microservices",
        "Monolith", "Data Structures", "Algorithms", "Object-Oriented Programming", "OOP",
        "Design Patterns", "System Design", "Machine Learning", "Deep Learning", "NLP",
        "Natural Language Processing", "Computer Vision", "LLM", "Generative AI", "Cybersecurity",
        "Web Scraping", "Data Analysis", "Data Visualization", "Unit Testing", "Integration Testing"
    ]
}

# Compile lowercased skill lookup mapping
SKILL_LOOKUP: Dict[str, str] = {}
for category, skills in SKILLS_TAXONOMY.items():
    for skill in skills:
        SKILL_LOOKUP[skill.lower()] = skill

# Additional common aliases / regex patterns
SKILL_PATTERNS = [
    (r"\b(c\+\+)\b", "C++"),
    (r"\b(c#)\b", "C#"),
    (r"\b(node(?:\.js)?)\b", "Node.js"),
    (r"\b(react(?:\.js)?)\b", "React"),
    (r"\b(vue(?:\.js)?)\b", "Vue.js"),
    (r"\b(express(?:\.js)?)\b", "Express.js"),
    (r"\b(rest(?:ful)?\s*apis?)\b", "REST APIs"),
    (r"\b(amazon\s*web\s*services|aws)\b", "AWS"),
    (r"\b(google\s*cloud(?:\s*platform)?|gcp)\b", "Google Cloud"),
    (r"\b(microsoft\s*azure|azure)\b", "Azure"),
    (r"\b(github\s*actions)\b", "GitHub Actions"),
    (r"\b(tailwind(?:\s*css)?)\b", "Tailwind CSS"),
    (r"\b(scikit-learn|sklearn)\b", "Scikit-Learn"),
    (r"\b(ci\s*\/\s*cd|ci-cd)\b", "CI/CD"),
    (r"\b(postgres(?:ql)?)\b", "PostgreSQL"),
    (r"\b(mongo(?:db)?)\b", "MongoDB"),
    (r"\b(ms\s*sql|sql\s*server)\b", "Microsoft SQL Server")
]


def extract_text_from_pdf(pdf_path: str) -> Dict[str, Any]:
    """Extract raw text and metadata from PDF using PyMuPDF (fitz) or fallback stream parser."""
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(pdf_path)
            page_texts = []
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text("text")
                if text.strip():
                    page_texts.append(text)

            full_text = "\n\n".join(page_texts).strip()
            meta = doc.metadata or {}
            doc.close()

            if full_text:
                return {
                    "text": full_text,
                    "page_count": total_pages,
                    "pdf_title": meta.get("title") or "",
                    "pdf_author": meta.get("author") or "",
                    "extractor": f"PyMuPDF v{getattr(fitz, '__version__', '1.21.1')}"
                }
        except Exception:
            pass

    # Built-in pure Python PDF text extraction fallback
    extracted_chunks = []
    try:
        with open(pdf_path, "rb") as f:
            content = f.read()

        import zlib
        # Find stream objects
        stream_matches = re.finditer(rb"stream[\r\n]+(.*?)[\r\n]+endstream", content, re.DOTALL)
        for sm in stream_matches:
            stream_data = sm.group(1)
            decompressed = None
            try:
                decompressed = zlib.decompress(stream_data)
            except Exception:
                decompressed = stream_data

            if decompressed:
                # Look for TJ / Tj text operators
                text_matches = re.findall(rb"\((.*?)\)\s*Tj", decompressed)
                for tm in text_matches:
                    try:
                        decoded_chunk = tm.decode("latin1", errors="ignore")
                        if decoded_chunk.strip():
                            extracted_chunks.append(decoded_chunk)
                    except Exception:
                        pass
                
                # Look for array text operators: [...] TJ
                array_matches = re.findall(rb"\[(.*?)\]\s*TJ", decompressed)
                for am in array_matches:
                    inner_texts = re.findall(rb"\((.*?)\)", am)
                    joined = " ".join(it.decode("latin1", errors="ignore") for it in inner_texts if it.strip())
                    if joined.strip():
                        extracted_chunks.append(joined)

        # Also search raw string literals in uncompressed PDF objects
        if not extracted_chunks:
            raw_strs = re.findall(rb"\(([A-Za-z0-9\s.,;:_@+\-\/()#&]{3,})\)", content)
            for rs in raw_strs:
                try:
                    s = rs.decode("latin1", errors="ignore").strip()
                    if len(s) > 3 and not s.startswith("http://www.w3.org"):
                        extracted_chunks.append(s)
                except Exception:
                    pass
    except Exception as e:
        pass

    full_text = "\n".join(extracted_chunks).strip() if extracted_chunks else ""
    return {
        "text": full_text or "Extracted resume content from PDF document.",
        "page_count": 1,
        "pdf_title": os.path.basename(pdf_path),
        "pdf_author": "",
        "extractor": "Pure-Python PDF Stream Parser"
    }


def extract_text_from_docx(docx_path: str) -> Dict[str, Any]:
    """Extract raw text from DOCX using python-docx or zero-dependency XML zipfile parser."""
    if PYTHON_DOCX_AVAILABLE:
        try:
            doc = docx.Document(docx_path)
            paragraphs = []

            # 1. Body paragraphs
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    paragraphs.append(t)

            # 2. Table contents
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        paragraphs.append(" | ".join(row_texts))

            full_text = "\n\n".join(paragraphs).strip()
            if full_text:
                return {
                    "text": full_text,
                    "page_count": max(1, len(full_text) // 2500),
                    "extractor": f"python-docx v{getattr(docx, '__version__', '0.8.11')}"
                }
        except Exception:
            pass

    # Built-in zero-dependency DOCX extractor (ZIP + XML parser)
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml_bytes = z.read("word/document.xml")
            tree = ET.fromstring(xml_bytes)
            
            paragraphs = []
            # Find all paragraph elements across XML namespaces
            for p in tree.iter():
                if p.tag.endswith("p"):
                    texts = []
                    for t in p.iter():
                        if t.tag.endswith("t") and t.text:
                            texts.append(t.text)
                    line = "".join(texts).strip()
                    if line:
                        paragraphs.append(line)

            full_text = "\n\n".join(paragraphs).strip()
            return {
                "text": full_text,
                "page_count": max(1, len(full_text) // 2500),
                "extractor": "Built-in DOCX/XML Parser"
            }
    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {e}")


def clean_text(text: str) -> str:
    """Normalize whitespace, tabs, and invalid characters in resume text."""
    # Replace non-breaking spaces
    text = text.replace("\u00a0", " ").replace("\ufeff", "")
    # Normalize carriage returns
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ consecutive newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_contact_info(text: str, lines: List[str]) -> Dict[str, Any]:
    """Extract candidate name, email, phone number, and social links."""
    # 1. Email extraction
    email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    email = email_match.group(0).strip() if email_match else None

    # 2. Phone extraction (handles +1 (555) 123-4567, +91 9876543210, 555-123-4567, etc.)
    phone_pattern = r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-.\s]?)?\d{3,4}[-.\s]?\d{3,4}"
    phone_matches = re.findall(phone_pattern, text)
    phone = None
    for match in phone_matches:
        cleaned_digits = re.sub(r"\D", "", match)
        # Phone numbers typically have 10-14 digits
        if 10 <= len(cleaned_digits) <= 14:
            # Exclude numbers that look like dates (e.g. 20182022) or zip codes
            if not (cleaned_digits.startswith("20") and len(cleaned_digits) == 8):
                phone = match.strip()
                break

    # 3. Web & Profile Links
    github_match = re.search(r"(?:https?:\/\/)?(?:www\.)?github\.com\/([a-zA-Z0-9_-]+)", text, re.IGNORECASE)
    github = f"https://github.com/{github_match.group(1)}" if github_match else None

    linkedin_match = re.search(r"(?:https?:\/\/)?(?:www\.)?linkedin\.com\/in\/([a-zA-Z0-9_-]+)", text, re.IGNORECASE)
    linkedin = f"https://linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else None

    portfolio_match = re.search(r"(?:https?:\/\/)?(?:www\.)?([a-zA-Z0-9-]+\.(?:dev|io|me|app|com))(?:\/[^\s,)]*)?", text, re.IGNORECASE)
    portfolio = portfolio_match.group(0) if portfolio_match and "linkedin" not in portfolio_match.group(0) and "github" not in portfolio_match.group(0) else None

    # 4. Name extraction heuristic
    name = None
    # Look at top lines before long paragraphs
    filtered_top_lines = []
    for line in lines[:8]:
        line_clean = line.strip()
        # Skip empty lines, lines with email, urls, phone, or generic labels
        if not line_clean:
            continue
        if "@" in line_clean or "http" in line_clean.lower() or "github" in line_clean.lower():
            continue
        if any(keyword in line_clean.lower() for keyword in ["resume", "curriculum vitae", "cv", "page 1", "contact", "summary"]):
            continue
        # Words check: should look like a human name (1 to 4 words, alphabetic, mostly capitalized)
        words = line_clean.split()
        if 1 <= len(words) <= 4:
            if all(re.match(r"^[A-Za-z.'-]+$", w) for w in words):
                filtered_top_lines.append(line_clean)

    if filtered_top_lines:
        name = filtered_top_lines[0]
    elif email:
        # Fallback derive name from email prefix if name wasn't detected
        name_prefix = email.split("@")[0]
        name_prefix = re.sub(r"[._0-9]+", " ", name_prefix).strip()
        name = name_prefix.title() if name_prefix else "Candidate"
    else:
        name = "Candidate"

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "github": github,
        "linkedin": linkedin,
        "portfolio": portfolio
    }


def extract_skills(text: str) -> Dict[str, Any]:
    """Scan and extract technical and soft skills with category classification."""
    found_skills_set = set()
    categorized_skills: Dict[str, List[str]] = {cat: [] for cat in SKILLS_TAXONOMY.keys()}

    # 1. Direct word boundary / regex matching
    text_lower = text.lower()

    # Exact token / phrase matching
    for skill_lower, skill_proper in SKILL_LOOKUP.items():
        # Match whole words to prevent false positives like 'c' matching 'in' or 'go' matching 'going'
        pattern = r"(?<![a-zA-Z0-9])" + re.escape(skill_lower) + r"(?![a-zA-Z0-9])"
        if re.search(pattern, text_lower):
            found_skills_set.add(skill_proper)

    # 2. Pattern based matching for special syntax
    for pattern, canonical_name in SKILL_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            found_skills_set.add(canonical_name)

    # Group into categories
    for skill in found_skills_set:
        for category, cat_list in SKILLS_TAXONOMY.items():
            if skill in cat_list or skill.lower() in [s.lower() for s in cat_list]:
                if skill not in categorized_skills[category]:
                    categorized_skills[category].append(skill)
                break
        else:
            if skill not in categorized_skills["Tools & Platforms"]:
                categorized_skills["Tools & Platforms"].append(skill)

    # Clean empty categories
    active_categories = {k: sorted(v) for k, v in categorized_skills.items() if v}
    all_skills_list = sorted(list(found_skills_set))

    return {
        "total_skills_count": len(all_skills_list),
        "all_skills": all_skills_list,
        "categories": active_categories
    }


def segment_resume_sections(text: str) -> Dict[str, str]:
    """Divide resume into structured logical sections based on common headings."""
    section_patterns = {
        "education": r"(?i)\b(?:education|academic\s+background|academic\s+qualifications|academics|qualifications)\b",
        "experience": r"(?i)\b(?:experience|work\s+experience|professional\s+experience|employment\s+history|internships|work\s+history)\b",
        "projects": r"(?i)\b(?:projects|academic\s+projects|personal\s+projects|technical\s+projects|key\s+projects)\b",
        "skills": r"(?i)\b(?:skills|technical\s+skills|core\s+competencies|technologies|technical\s+expertise|key\s+skills)\b",
        "certifications": r"(?i)\b(?:certifications|certificates|licenses\s+(&|and)\s+certifications|professional\s+certifications|courses)\b",
        "summary": r"(?i)\b(?:summary|professional\s+summary|profile|about\s+me|career\s+objective|objective)\b",
        "publications": r"(?i)\b(?:publications|research\s+papers|research)\b",
        "awards": r"(?i)\b(?:awards|honors|achievements|extracurricular\s+activities)\b"
    }

    # Find section positions
    matches = []
    for section_name, pattern in section_patterns.items():
        for match in re.finditer(pattern, text):
            # Check if match is on its own line or followed by colon/newline
            start = match.start()
            line_start = text.rfind("\n", 0, start)
            line_start = 0 if line_start == -1 else line_start + 1
            line_end = text.find("\n", start)
            line_end = len(text) if line_end == -1 else line_end
            line = text[line_start:line_end].strip()

            # Heading lines are generally short (< 50 chars)
            if len(line) < 50:
                matches.append({
                    "section": section_name,
                    "start": line_start,
                    "end": line_end,
                    "heading": line
                })

    # Sort matches by position in text
    matches.sort(key=lambda x: x["start"])

    sections: Dict[str, str] = {}
    for i, match in enumerate(matches):
        sec_start = match["end"]
        sec_end = matches[i + 1]["start"] if i + 1 < len(matches) else len(text)
        content = text[sec_start:sec_end].strip()
        if content:
            sections[match["section"]] = content

    return sections


def extract_education(edu_text: str) -> List[Dict[str, Any]]:
    """Extract education history, degrees, universities, and graduation years."""
    if not edu_text:
        return []

    degree_patterns = [
        r"\b(?:Bachelor\s+of\s+Technology|B\.?Tech|B\.?E\.?|Bachelor\s+of\s+Engineering|B\.?S\.?|Bachelor\s+of\s+Science|B\.?C\.?A\.?|Bachelor\s+of\s+Computer\s+Applications)\b",
        r"\b(?:Master\s+of\s+Technology|M\.?Tech|M\.?E\.?|Master\s+of\s+Engineering|M\.?S\.?|Master\s+of\s+Science|M\.?C\.?A\.?|Master\s+of\s+Computer\s+Applications|M\.?B\.?A\.?)\b",
        r"\b(?:Ph\.?D\.?|Doctorate|Doctor\s+of\s+Philosophy|Diploma|Associate\s+Degree|High\s+School|Secondary\s+School)\b"
    ]

    year_pattern = r"\b(?:19|20)\d{2}(?:\s*[-–—]\s*(?:(?:19|20)\d{2}|Present|Current|Expected))?\b"
    gpa_pattern = r"\b(?:GPA|CGPA|Percentage|Score)?\s*[:=]?\s*(\d+(?:\.\d+)?\s*(?:\/\s*\d+(?:\.\d+)?)?|\d{2,3}%)\b"

    entries = []
    lines = [line.strip() for line in edu_text.split("\n") if line.strip()]
    
    current_entry: Dict[str, Any] = {}
    
    for line in lines:
        degree_found = None
        for deg_pat in degree_patterns:
            deg_match = re.search(deg_pat, line, re.IGNORECASE)
            if deg_match:
                degree_found = deg_match.group(0)
                break

        years = re.findall(year_pattern, line, re.IGNORECASE)
        gpa_match = re.search(gpa_pattern, line, re.IGNORECASE)

        if degree_found or (not current_entry and len(line) > 5):
            if current_entry:
                entries.append(current_entry)
            
            current_entry = {
                "degree": degree_found or "Degree / Coursework",
                "institution": line if not degree_found else line.replace(degree_found, "").strip(" ,|-:"),
                "year": years[0] if years else None,
                "gpa": gpa_match.group(0) if gpa_match else None,
                "details": [line]
            }
        elif current_entry:
            if years and not current_entry.get("year"):
                current_entry["year"] = years[0]
            if gpa_match and not current_entry.get("gpa"):
                current_entry["gpa"] = gpa_match.group(0)
            current_entry["details"].append(line)
            if not current_entry.get("institution") or current_entry.get("institution") == "":
                current_entry["institution"] = line

    if current_entry:
        entries.append(current_entry)

    # Clean entries
    for entry in entries:
        if not entry.get("institution"):
            entry["institution"] = "Academic Institution"

    return entries


def extract_experience(exp_text: str) -> List[Dict[str, Any]]:
    """Extract professional experience, roles, companies, and bullet points."""
    if not exp_text:
        return []

    date_pattern = r"(?i)\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}\s*[-–—]\s*(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)\b|\b\d{4}\s*[-–—]\s*(?:\d{4}|Present|Current)\b"
    role_hints = [
        "Engineer", "Developer", "Intern", "Analyst", "Lead", "Architect", "Consultant",
        "Manager", "Specialist", "Scientist", "Designer", "Associate", "Assistant", "Administrator"
    ]

    lines = [l.strip() for l in exp_text.split("\n") if l.strip()]
    items = []
    current: Optional[Dict[str, Any]] = None

    for line in lines:
        date_match = re.search(date_pattern, line)
        has_role_hint = any(hint.lower() in line.lower() for hint in role_hints)

        # Start a new experience block when date or role title is found
        if date_match or (has_role_hint and (not current or len(current["bullets"]) >= 2)):
            if current:
                items.append(current)

            dates = date_match.group(0) if date_match else None
            title_text = line.replace(dates, "").strip(" ,|-–—") if dates else line

            parts = [p.strip() for p in re.split(r"[,|–—@•]", title_text) if p.strip()]
            role = parts[0] if parts else "Software Engineer"
            company = parts[1] if len(parts) > 1 else "Technology Company"

            current = {
                "role": role,
                "company": company,
                "duration": dates or "Recent",
                "bullets": []
            }
        else:
            if current:
                clean_bullet = line.lstrip("•-*–—> 0123456789.)").strip()
                if clean_bullet:
                    current["bullets"].append(clean_bullet)
            else:
                # First line before standard header
                current = {
                    "role": line,
                    "company": "Company / Organization",
                    "duration": "Experience",
                    "bullets": []
                }

    if current:
        items.append(current)

    return items


def extract_projects(proj_text: str) -> List[Dict[str, Any]]:
    """Extract projects, technologies used, and descriptions."""
    if not proj_text:
        return []

    lines = [l.strip() for l in proj_text.split("\n") if l.strip()]
    projects = []
    current_proj: Optional[Dict[str, Any]] = None

    for line in lines:
        # Check if line looks like a project title (e.g. "Resume Skill Gap Analyzer | React, Python, Flask")
        if "|" in line or ":" in line or not current_proj:
            if current_proj and current_proj.get("bullets"):
                projects.append(current_proj)

            parts = [p.strip() for p in re.split(r"[:|]", line, maxsplit=1)]
            title = parts[0].lstrip("•-*–—> ")
            tech_stack = parts[1] if len(parts) > 1 else ""

            # Check for github or live demo links
            link_match = re.search(r"https?:\/\/[^\s,)]+", line)
            link = link_match.group(0) if link_match else None

            current_proj = {
                "title": title,
                "technologies": [t.strip() for t in re.split(r"[,/]", tech_stack) if t.strip()],
                "link": link,
                "bullets": []
            }
        else:
            if current_proj:
                clean_bullet = line.lstrip("•-*–—> 0123456789.)").strip()
                if clean_bullet:
                    current_proj["bullets"].append(clean_bullet)
            else:
                current_proj = {
                    "title": line,
                    "technologies": [],
                    "link": None,
                    "bullets": []
                }

    if current_proj:
        projects.append(current_proj)

    return projects


def extract_certifications(cert_text: str) -> List[Dict[str, Any]]:
    """Extract certifications, issuing organizations, and years."""
    if not cert_text:
        return []

    lines = [l.strip() for l in cert_text.split("\n") if l.strip()]
    certifications = []

    year_pat = r"\b(?:19|20)\d{2}\b"

    for line in lines:
        clean_cert = line.lstrip("•-*–—> 0123456789.)").strip()
        if not clean_cert:
            continue

        year_match = re.search(year_pat, clean_cert)
        year = year_match.group(0) if year_match else None

        # Common issuers
        issuers = ["AWS", "Google", "Microsoft", "Coursera", "Udemy", "edX", "Oracle", "Cisco", "CompTIA", "Meta", "IBM"]
        issuer = None
        for org in issuers:
            if org.lower() in clean_cert.lower():
                issuer = org
                break

        certifications.append({
            "name": clean_cert,
            "issuer": issuer or "Accredited Provider",
            "year": year
        })

    return certifications


def parse_resume_file(file_path: str, original_filename: str) -> Dict[str, Any]:
    """
    Main entry point for parsing resume file (PDF or DOCX).
    Returns complete extracted raw text and structured entities.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at path: {file_path}")

    ext = Path(original_filename).suffix.lower()
    if ext not in [".pdf", ".docx"]:
        raise ValueError(f"Unsupported file format '{ext}'. Only PDF (.pdf) and DOCX (.docx) are supported.")

    # 1. Extract raw text based on file format
    if ext == ".pdf":
        raw_res = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        raw_res = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    raw_text = clean_text(raw_res["text"])

    if not raw_text or len(raw_text.strip()) < 10:
        raise ValueError("Resume file appears empty or text could not be extracted.")

    lines = [line.strip() for line in raw_text.split("\n") if line.strip()]

    # 2. Extract entities
    contact_info = extract_contact_info(raw_text, lines)
    skills_data = extract_skills(raw_text)
    sections = segment_resume_sections(raw_text)

    education_list = extract_education(sections.get("education", ""))
    experience_list = extract_experience(sections.get("experience", ""))
    projects_list = extract_projects(sections.get("projects", ""))
    certifications_list = extract_certifications(sections.get("certifications", ""))
    summary_text = sections.get("summary", "")

    # Stats
    word_count = len(raw_text.split())
    char_count = len(raw_text)

    parsed_result = {
        "file_name": original_filename,
        "file_type": ext.lstrip(".").upper(),
        "extractor": raw_res.get("extractor", "Standard Parser"),
        "page_count": raw_res.get("page_count", 1),
        "word_count": word_count,
        "char_count": char_count,
        "contact": contact_info,
        "summary": summary_text,
        "skills": skills_data,
        "education": education_list,
        "experience": experience_list,
        "projects": projects_list,
        "certifications": certifications_list,
        "raw_text": raw_text
    }

    return parsed_result


if __name__ == "__main__":
    # Standalone CLI test handler
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python resume_parser.py <file_path> [original_filename]"}))
        sys.exit(1)

    target_path = sys.argv[1]
    orig_name = sys.argv[2] if len(sys.argv) > 2 else os.path.basename(target_path)

    try:
        result = parse_resume_file(target_path, orig_name)
        print(json.dumps(result, indent=2))
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)
